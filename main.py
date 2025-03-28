from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from pdf2image import convert_from_path
from PIL import Image, ImageFilter, ImageEnhance, ImageOps
import os
import sys
import json
import logging
import traceback
import numpy as np
import PyPDF2
import argparse

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Try to import advanced libraries
try:
    import pdfplumber
    import pikepdf
    import pytesseract
    import cv2
    import pandas as pd
    import docx
    import tabula
    import camelot
    import nltk
    import spacy
    ADVANCED_IMPORTS_AVAILABLE = True
except ImportError as e:
    logger.warning(f"Some advanced imports are not available: {e}")
    ADVANCED_IMPORTS_AVAILABLE = False

# Add a function to download spaCy models if needed
def download_spacy_model(model_name="en_core_web_sm"):
    """Download spaCy model if not available."""
    try:
        # Check if spacy is available
        import spacy
        try:
            # Try to load the model to see if it's installed
            spacy.load(model_name)
            logger.info(f"spaCy model '{model_name}' is already available.")
            return True
        except OSError:
            # Model is not installed, so download it
            logger.info(f"Downloading spaCy model '{model_name}'...")
            import subprocess
            import sys
            subprocess.check_call([sys.executable, "-m", "spacy", "download", model_name])
            logger.info(f"Successfully downloaded spaCy model '{model_name}'")
            return True
    except Exception as e:
        logger.warning(f"Failed to download spaCy model: {e}")
        return False

# Add a function to locate Tesseract executable
def find_tesseract_executable():
    """Try to find Tesseract executable in common installation paths."""
    import os
    import platform
    
    # Common installation paths for different operating systems
    if platform.system() == 'Windows':
        common_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Tesseract-OCR\tesseract.exe',
            r'D:\Tesseract-OCR\tesseract.exe',
            r'E:\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files\Tesseract\tesseract.exe',
            # Add more common Windows installation paths here
        ]
    elif platform.system() == 'Darwin':  # macOS
        common_paths = [
            '/usr/local/bin/tesseract',
            '/opt/homebrew/bin/tesseract',
            '/usr/bin/tesseract',
            # Add more common macOS installation paths here
        ]
    else:  # Linux and other Unix-like systems
        common_paths = [
            '/usr/bin/tesseract',
            '/usr/local/bin/tesseract',
            '/opt/tesseract/bin/tesseract',
            # Add more common Linux installation paths here
        ]
    
    # Check if the executable exists in common paths
    for path in common_paths:
        if os.path.isfile(path):
            logger.info(f"Found Tesseract executable at: {path}")
            return path
            
    # If not found in common paths, check in system PATH
    try:
        import shutil
        tesseract_path = shutil.which('tesseract')
        if tesseract_path:
            logger.info(f"Found Tesseract executable in PATH: {tesseract_path}")
            return tesseract_path
    except Exception as e:
        logger.warning(f"Error checking for Tesseract in PATH: {e}")
        
    logger.warning("Tesseract executable not found in common paths or PATH")
    return None

class DataExtractor:
    """Enhanced class for extracting data from various file formats with advanced processing."""
    
    def __init__(self, ocr_lang='eng', extract_tables=True, output_dir=None, 
                 dpi=300, preprocess=True, preserve_layout=True, ocr_mode="advanced",
                 font_analysis=True, extract_forms=True, max_threads=4, tesseract_path=None):
        """
        Initialize the DataExtractor with advanced options.
        
        Args:
            ocr_lang (str or list): Language(s) for OCR processing
            extract_tables (bool): Whether to extract tables from documents
            output_dir (str): Directory to save extracted data
            dpi (int): DPI for image conversion (higher = better quality but slower)
            preprocess (bool): Apply image preprocessing for better OCR
            preserve_layout (bool): Attempt to preserve document layout
            ocr_mode (str): OCR mode - "basic", "advanced", or "deep"
            font_analysis (bool): Perform font analysis on PDF documents
            extract_forms (bool): Extract form fields from documents
            max_threads (int): Maximum number of threads for parallel processing
            tesseract_path (str): Path to the Tesseract executable
        """
        # Convert language to list for multi-language support
        self.ocr_lang = ocr_lang if isinstance(ocr_lang, list) else [ocr_lang]
        self.extract_tables = extract_tables
        self.output_dir = output_dir or os.path.join(os.getcwd(), "extracted_data")
        self.dpi = dpi
        self.preprocess = preprocess
        self.preserve_layout = preserve_layout
        self.ocr_mode = ocr_mode
        self.font_analysis = font_analysis
        self.extract_forms = extract_forms
        self.max_threads = max_threads
        self.tesseract_path = tesseract_path
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        # Initialize format-specific extractors
        self._initialize_extractors()
    
    def _initialize_extractors(self):
        """Initialize format-specific extractors with required libraries."""
        # Attempt to import required libraries
        try:
            # For PDFs
            import PyPDF2
            self.pdf_extractor = PyPDF2.PdfReader
            logger.info("Initialized PDF extractor with PyPDF2")
            
            # For advanced PDF processing
            try:
                import pdfplumber
                self.pdfplumber_available = True
                logger.info("Enhanced PDF processing enabled with pdfplumber")
            except ImportError:
                self.pdfplumber_available = False
                logger.warning("pdfplumber not available. Advanced PDF layout analysis will be limited.")

            # For PDF/A conversion and font analysis
            try:
                import pikepdf
                self.pikepdf_available = True
                logger.info("Enhanced PDF manipulation enabled with pikepdf")
            except ImportError:
                self.pikepdf_available = False
                logger.warning("pikepdf not available. Some advanced PDF features will be limited.")
            
            # For OCR on PDFs and images
            try:
                import pytesseract
                from pdf2image import convert_from_path
                from PIL import Image, ImageFilter, ImageEnhance, ImageOps
                
                # Configure pytesseract with the tesseract path if provided
                if self.tesseract_path:
                    pytesseract.pytesseract.tesseract_cmd = self.tesseract_path
                    logger.info(f"Using Tesseract executable at: {self.tesseract_path}")
                elif hasattr(pytesseract.pytesseract, 'tesseract_cmd') and os.path.isfile(pytesseract.pytesseract.tesseract_cmd):
                    logger.info(f"Using existing Tesseract path: {pytesseract.pytesseract.tesseract_cmd}")
                else:
                    # Try to find Tesseract executable in common paths
                    found_path = find_tesseract_executable()
                    if found_path:
                        pytesseract.pytesseract.tesseract_cmd = found_path
                        logger.info(f"Automatically set Tesseract path to: {found_path}")
                    else:
                        logger.warning("Could not locate Tesseract executable. Please specify the path manually.")
                
                # Test if Tesseract is working
                try:
                    pytesseract.get_tesseract_version()
                    logger.info(f"Tesseract version: {pytesseract.get_tesseract_version()}")
                    self.ocr_available = True
                except Exception as e:
                    logger.warning(f"Tesseract executable not functional: {e}")
                    logger.warning("OCR features will be limited or unavailable.")
                    self.ocr_available = False
                
                self.pytesseract = pytesseract
                self.pdf2image = convert_from_path
                self.Image = Image
                self.ImageFilter = ImageFilter
                self.ImageEnhance = ImageEnhance
                self.ImageOps = ImageOps
                
                if self.ocr_available:
                    logger.info("OCR capabilities enabled with pytesseract")
                
                # Set up OCR configurations based on mode
                self.ocr_config = ""
                if self.ocr_mode == "advanced":
                    self.ocr_config = "--oem 1 --psm 6"
                elif self.ocr_mode == "deep":
                    self.ocr_config = "--oem 1 --psm 11 -c tessedit_do_invert=0 -c textord_tablefind_recognize_tables=1"
                
                # For advanced OCR and image processing
                try:
                    import cv2
                    self.cv2_available = True
                    self.cv2 = cv2
                    logger.info("Advanced image processing enabled with OpenCV")
                except ImportError:
                    self.cv2_available = False
                    logger.warning("OpenCV not available. Advanced image preprocessing will be limited.")
            except ImportError:
                logger.warning("OCR libraries not available. Install pytesseract and pdf2image for OCR support.")
                self.ocr_available = False
            
            # For Excel files
            try:
                import openpyxl
                self.excel_available = True
                logger.info("Excel extraction enabled with pandas/openpyxl")
            except ImportError:
                logger.warning("Excel library not available. Install openpyxl for better Excel support.")
                self.excel_available = True  # pandas can still handle Excel
            
            # For Word documents
            try:
                import docx
                self.docx = docx
                self.word_available = True
                logger.info("Word document extraction enabled")
            except ImportError:
                logger.warning("Word library not available. Install python-docx for Word document support.")
                self.word_available = False
            
            # For advanced table extraction
            try:
                import tabula
                import camelot
                self.tabula = tabula
                self.camelot = camelot
                self.advanced_table_extraction = True
                logger.info("Advanced table extraction enabled with tabula and camelot")
            except ImportError:
                self.advanced_table_extraction = False
                logger.warning("Advanced table extraction libraries not available. Install tabula-py and camelot-py for better table extraction.")

            # For natural language processing and structure identification
            try:
                import nltk
                import spacy
                self.nltk_available = True
                self.spacy_available = True
                
                # Download NLTK resources if not already downloaded
                try:
                    nltk.data.find('tokenizers/punkt')
                except LookupError:
                    nltk.download('punkt', quiet=True)
                
                try:
                    nltk.data.find('averaged_perceptron_tagger')
                except LookupError:
                    nltk.download('averaged_perceptron_tagger', quiet=True)
                
                # Try to load spaCy model, download if needed, or use fallback
                spacy_model = "en_core_web_sm"
                try:
                    # Try to load the model
                    self.nlp = spacy.load(spacy_model)
                    logger.info(f"Loaded spaCy model '{spacy_model}'")
                except OSError:
                    # Model not found, try to download it
                    if download_spacy_model(spacy_model):
                        try:
                            self.nlp = spacy.load(spacy_model)
                            logger.info(f"Successfully loaded spaCy model '{spacy_model}' after downloading")
                        except Exception as e:
                            # If still can't load, create a blank model as fallback
                            logger.warning(f"Could not load spaCy model after downloading: {e}")
                            self.nlp = spacy.blank("en")
                            logger.info("Using blank spaCy model as fallback")
                    else:
                        # If download failed, use blank model
                        self.nlp = spacy.blank("en")
                        logger.info("Using blank spaCy model as fallback")
                
                logger.info("NLP capabilities enabled for structure detection")
            except ImportError:
                self.nltk_available = False
                self.spacy_available = False
                logger.warning("NLP libraries not available. Install nltk and spacy for advanced text structure analysis.")
                
        except ImportError as e:
            logger.error(f"Failed to initialize extractors: {e}")
            logger.info("Please install required packages using: pip install PyPDF2 pandas openpyxl pytesseract pdf2image python-docx Pillow pdfplumber pikepdf tabula-py camelot-py opencv-python nltk spacy")
            sys.exit(1)
    
    def extract(self, file_path):
        """
        Extract data from a file based on its extension.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: Extracted data with metadata
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext in ['.pdf']:
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.xlsx', '.xls', '.csv']:
                return self._extract_from_excel(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_word(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
                return self._extract_from_image(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_ext}")
                return {"error": f"Unsupported file format: {file_ext}"}
        except Exception as e:
            logger.error(f"Error extracting data from {file_path}: {e}")
            traceback.print_exc()
            return {"error": str(e)}
    
    def _extract_from_pdf(self, file_path):
        """Extract text, tables, and structure from PDF files with advanced processing."""
        logger.info(f"Extracting data from PDF: {file_path}")
        result = {
            "file_path": file_path, 
            "type": "pdf", 
            "text": [], 
            "tables": [],
            "metadata": {},
            "structure": {},
            "fonts": [],
            "forms": []
        }
        
        # Extract metadata with pikepdf if available
        if self.pikepdf_available:
            try:
                import pikepdf
                pdf = pikepdf.Pdf.open(file_path)
                docinfo = pdf.docinfo
                if docinfo:
                    for key, value in docinfo.items():
                        if str(key).startswith('/'):
                            clean_key = str(key)[1:]  # Remove leading slash
                            result["metadata"][clean_key] = str(value)
                logger.info(f"Extracted PDF metadata: {len(result['metadata'])} fields")
            except Exception as e:
                logger.warning(f"Error extracting PDF metadata: {e}")
        
        # Advanced PDF processing with pdfplumber
        if self.pdfplumber_available:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    # Extract font information if requested
                    if self.font_analysis:
                        fonts_by_page = {}
                        for i, page in enumerate(pdf.pages):
                            fonts = {}
                            page_fonts = {}
                            
                            # Extract font information from page
                            if hasattr(page, '_objs') and page._objs:
                                for obj in page._objs:
                                    if hasattr(obj, 'get') and obj.get('Font'):
                                        fonts.update(obj.get('Font', {}))
                            
                            # Process font details for better analysis
                            if fonts:
                                for font_id, font_dict in fonts.items():
                                    if hasattr(font_dict, 'resolve'):
                                        font_obj = font_dict.resolve()
                                        if hasattr(font_obj, 'get'):
                                            # Extract the base font, type, and encoding
                                            base_font = font_obj.get('BaseFont', 'Unknown')
                                            font_type = font_obj.get('Subtype', 'Unknown')
                                            encoding = font_obj.get('Encoding', 'Unknown')
                                            
                                            page_fonts[str(font_id)] = {
                                                'base_font': str(base_font),
                                                'type': str(font_type),
                                                'encoding': str(encoding)
                                            }
                            
                            # Get character-level font usage for detailed analysis
                            if page.chars:
                                # Group by font to analyze text with specific fonts
                                font_usage = {}
                                for char in page.chars:
                                    font_name = char.get('fontname', 'Unknown')
                                    if font_name not in font_usage:
                                        font_usage[font_name] = {
                                            'count': 0,
                                            'sizes': set(),
                                            'sample': ''
                                        }
                                    
                                    font_usage[font_name]['count'] += 1
                                    if len(font_usage[font_name]['sample']) < 20:
                                        font_usage[font_name]['sample'] += char.get('text', '')
                                    font_usage[font_name]['sizes'].add(round(char.get('size', 0), 1))
                                
                                # Convert sets to lists for JSON serialization
                                for font, data in font_usage.items():
                                    data['sizes'] = sorted(list(data['sizes']))
                                
                                page_fonts['usage'] = font_usage
                            
                            fonts_by_page[str(i+1)] = page_fonts
                        
                        result["fonts"] = fonts_by_page
                        logger.info(f"Identified {len(fonts_by_page)} pages with font information")
                    
                    # Extract form fields if requested
                    if self.extract_forms:
                        form_fields = []
                        for i, page in enumerate(pdf.pages):
                            page_fields = []
                            annotations = page.annots if hasattr(page, 'annots') else []
                            
                            if annotations:
                                for j, annot in enumerate(annotations):
                                    if annot and annot.get('subtype') == 'Widget':
                                        field = {
                                            'page': i+1,
                                            'id': f"field_{i}_{j}",
                                            'type': annot.get('ft', 'Unknown'),
                                            'name': annot.get('t', f"Field_{j}"),
                                            'value': annot.get('v', ''),
                                            'rect': annot.get('rect', [0, 0, 0, 0]),
                                            'required': annot.get('ff', 0) & 2 == 2,
                                            'readonly': annot.get('ff', 0) & 1 == 1,
                                        }
                                        page_fields.append(field)
                            
                            if page_fields:
                                form_fields.extend(page_fields)
                        
                        result["forms"] = form_fields
                        if form_fields:
                            logger.info(f"Extracted {len(form_fields)} form fields from the document")
                    
                    # Extract text with layout analysis if requested
                    if self.preserve_layout:
                        with ThreadPoolExecutor(max_workers=self.max_threads) as executor:
                            # Process pages in parallel
                            future_to_page = {executor.submit(self._process_page_with_layout, 
                                                              page, i): (i, page) 
                                            for i, page in enumerate(pdf.pages)}
                            
                            # Collect results as they complete
                            for future in as_completed(future_to_page):
                                i, extracted_text = future.result()
                                if extracted_text:
                                    result["text"].append({
                                        "page": i + 1,
                                        "content": extracted_text,
                                    })
                    else:
                        # Simple text extraction
                        for i, page in enumerate(pdf.pages):
                            text = page.extract_text(x_tolerance=3, y_tolerance=3)
                            if text and text.strip():
                                result["text"].append({
                                    "page": i + 1,
                                    "content": text,
                                })
            except Exception as e:
                logger.warning(f"Error in pdfplumber extraction: {e}. Will try other methods.")
        
        # If no text was extracted or we need to use PyPDF2 as fallback
        if not result["text"]:
            try:
                reader = self.pdf_extractor(file_path)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        result["text"].append({
                            "page": page_num + 1,
                            "content": text
                        })
            except Exception as e:
                logger.warning(f"Error in PDF text extraction: {e}. Will try OCR if available.")
        
        # If still no text was extracted or OCR is needed for scanned documents
        if (not result["text"] or self.ocr_mode == "deep") and self.ocr_available:
            logger.info(f"Using advanced OCR processing on PDF: {file_path}")
            try:
                # Convert PDF to images
                images = self.pdf2image(file_path, dpi=self.dpi)
                
                # Process pages in parallel
                with ThreadPoolExecutor(max_workers=self.max_threads) as executor:
                    future_to_page = {executor.submit(self._process_image_for_ocr, 
                                                     image, i): (i, image) 
                                     for i, image in enumerate(images)}
                    
                    # Collect results as they complete
                    for future in as_completed(future_to_page):
                        i, ocr_result = future.result()
                        result["text"].append({
                            "page": i + 1,
                            "content": ocr_result,
                        })
                
                logger.info(f"OCR processing completed for {len(images)} pages")
            except Exception as e:
                logger.error(f"Advanced OCR extraction failed: {e}")
        
        # Extract and analyze document structure
        if result["text"] and (self.nltk_available or self.spacy_available):
            try:
                structure = self._analyze_document_structure([page["content"] for page in result["text"]])
                result["structure"] = structure
                logger.info("Completed document structure analysis")
            except Exception as e:
                logger.warning(f"Error analyzing document structure: {e}")
        
        # Extract tables with advanced analysis if requested
        if self.extract_tables:
            try:
                if self.advanced_table_extraction:
                    # Try camelot first for more accurate table extraction
                    try:
                        tables = self.camelot.read_pdf(file_path, pages='all', flavor='lattice')
                        if len(tables) > 0:
                            for i, table in enumerate(tables):
                                df = table.df
                                result["tables"].append({
                                    "table_id": i + 1,
                                    "page": table.page,
                                    "accuracy": table.accuracy,
                                    "rows": len(df),
                                    "cols": len(df.columns),
                                    "data": df.to_dict('records')
                                })
                            logger.info(f"Extracted {len(tables)} tables using camelot (lattice mode)")
                        else:
                            # Try stream mode as fallback
                            tables = self.camelot.read_pdf(file_path, pages='all', flavor='stream')
                            for i, table in enumerate(tables):
                                df = table.df
                                result["tables"].append({
                                    "table_id": i + 1,
                                    "page": table.page,
                                    "accuracy": table.accuracy,
                                    "rows": len(df),
                                    "cols": len(df.columns),
                                    "data": df.to_dict('records')
                                })
                            logger.info(f"Extracted {len(tables)} tables using camelot (stream mode)")
                    except Exception as e:
                        logger.warning(f"Camelot table extraction error: {e}. Trying tabula...")
                        
                    # Use tabula as fallback
                    if not result["tables"]:
                        tables = self.tabula.read_pdf(file_path, pages='all', multiple_tables=True)
                        for i, table in enumerate(tables):
                            result["tables"].append({
                                "table_id": i + 1,
                                "rows": len(table),
                                "cols": len(table.columns),
                                "data": table.to_dict('records')
                            })
                        logger.info(f"Extracted {len(tables)} tables using tabula")
                else:
                    # Fallback to tabula-only extraction
                    import tabula
                    tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
                    for i, table in enumerate(tables):
                        result["tables"].append({
                            "table_id": i + 1,
                            "data": table.to_dict('records')
                        })
            except ImportError:
                logger.warning("Table extraction libraries not available. Install tabula-py or camelot-py.")
            except Exception as e:
                logger.error(f"Error extracting tables: {e}")
        
        # Save extracted text to file
        output_text = self._format_extracted_text(result)
        output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_extracted.txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output_text)
        
        # Save structured data to JSON
        json_output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_structured.json")
        structured_data = {
            "metadata": result["metadata"],
            "structure": result["structure"],
            "tables": result["tables"],
            "fonts": result["fonts"],
            "forms": result["forms"]
        }
        
        with open(json_output_path, 'w', encoding='utf-8') as f:
            json.dump(structured_data, f, indent=2, ensure_ascii=False)
        
        result["output_file"] = output_path
        result["structured_data_file"] = json_output_path
        return result
    
    def _process_page_with_layout(self, page, page_num):
        """Process a page with layout preservation."""
        try:
            # Extract text in small blocks to better preserve layout
            blocks = []
            
            # Extract text by character with positions for better layout preservation
            chars = page.chars
            if chars:
                # Group chars by line
                current_line = []
                current_y = None
                y_tolerance = 3  # Adjust based on document spacing
                
                # Sort chars by vertical position (y) and then by horizontal position (x)
                sorted_chars = sorted(chars, key=lambda c: (c['top'], c['x0']))
                
                for char in sorted_chars:
                    if current_y is None:
                        current_y = char['top']
                        current_line.append(char)
                    elif abs(char['top'] - current_y) <= y_tolerance:
                        current_line.append(char)
                    else:
                        # Sort current line by x position and join
                        sorted_line = sorted(current_line, key=lambda c: c['x0'])
                        line_text = ''.join(c['text'] for c in sorted_line)
                        blocks.append(line_text)
                        
                        # Start new line
                        current_line = [char]
                        current_y = char['top']
                
                # Don't forget the last line
                if current_line:
                    sorted_line = sorted(current_line, key=lambda c: c['x0'])
                    line_text = ''.join(c['text'] for c in sorted_line)
                    blocks.append(line_text)
                
                return page_num, '\n'.join(blocks)
            
            # Fallback to standard text extraction
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            return page_num, text
        except Exception as e:
            logger.warning(f"Error processing page {page_num+1} layout: {e}")
            # Fall back to standard extraction
            try:
                return page_num, page.extract_text()
            except:
                return page_num, ""
    
    def _process_image_for_ocr(self, image, page_num):
        """Process an image for optimal OCR results with advanced preprocessing."""
        try:
            if self.preprocess and self.cv2_available:
                # Convert PIL image to OpenCV format
                img = np.array(image)
                img = img[:, :, ::-1].copy()  # RGB to BGR
                
                # Apply a series of preprocessing steps
                # 1. Convert to grayscale
                gray = self.cv2.cvtColor(img, self.cv2.COLOR_BGR2GRAY)
                
                # 2. Apply adaptive thresholding
                binary = self.cv2.adaptiveThreshold(
                    gray, 255, self.cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                    self.cv2.THRESH_BINARY, 11, 2
                )
                
                # 3. Apply morphological operations to clean up noise
                kernel = np.ones((1, 1), np.uint8)
                opening = self.cv2.morphologyEx(binary, self.cv2.MORPH_OPEN, kernel)
                
                # 4. Apply bilateral filter to preserve edges while removing noise
                denoised = self.cv2.bilateralFilter(opening, 9, 75, 75)
                
                # Convert back to PIL image for Tesseract
                processed_image = self.Image.fromarray(denoised)
                
                # Apply additional PIL-based enhancements
                processed_image = self.ImageOps.autocontrast(processed_image)
                enhancer = self.ImageEnhance.Contrast(processed_image)
                processed_image = enhancer.enhance(2.0)
                
                # Resize image for better OCR if it's too small
                width, height = processed_image.size
                if width < 1000 or height < 1000:
                    scale_factor = max(1000 / width, 1000 / height)
                    new_width = int(width * scale_factor)
                    new_height = int(height * scale_factor)
                    processed_image = processed_image.resize((new_width, new_height), self.Image.LANCZOS)
                
                # Apply OCR with advanced configuration
                if self.ocr_mode == "deep":
                    # Use multiple languages and special configuration for deep mode
                    lang_str = "+".join(self.ocr_lang)
                    ocr_result = self.pytesseract.image_to_string(
                        processed_image,
                        lang=lang_str,
                        config=self.ocr_config
                    )
                    
                    # For deep mode, also try to get structural information
                    try:
                        # Use image_to_data to get position and confidence of each word
                        ocr_data = self.pytesseract.image_to_data(
                            processed_image, 
                            lang=lang_str,
                            config=self.ocr_config,
                            output_type=self.pytesseract.Output.DICT
                        )
                        
                        # Organize data into a structured format
                        structured_lines = []
                        current_line = ""
                        current_line_num = -1
                        
                        for i in range(len(ocr_data["text"])):
                            if ocr_data["text"][i].strip():
                                # New line detected
                                if ocr_data["line_num"][i] != current_line_num:
                                    if current_line:
                                        structured_lines.append(current_line.strip())
                                    current_line = ocr_data["text"][i]
                                    current_line_num = ocr_data["line_num"][i]
                                else:
                                    current_line += " " + ocr_data["text"][i]
                        
                        # Add the last line
                        if current_line:
                            structured_lines.append(current_line.strip())
                        
                        ocr_result = "\n".join(structured_lines)
                    except Exception as e:
                        logger.warning(f"Error in structured OCR for page {page_num+1}: {e}")
                else:
                    # Basic OCR with multiple language support
                    lang_str = "+".join(self.ocr_lang)
                    ocr_result = self.pytesseract.image_to_string(
                        processed_image,
                        lang=lang_str,
                        config=self.ocr_config
                    )
                
                return page_num, ocr_result
            else:
                # Basic OCR without preprocessing
                lang_str = "+".join(self.ocr_lang)
                text = self.pytesseract.image_to_string(image, lang=lang_str, config=self.ocr_config)
                return page_num, text
                
        except Exception as e:
            logger.warning(f"Error preprocessing image for OCR on page {page_num+1}: {e}")
            # Fall back to basic OCR
            try:
                lang_str = "+".join(self.ocr_lang)
                text = self.pytesseract.image_to_string(image, lang=lang_str)
                return page_num, text
            except:
                return page_num, ""
    
    def _analyze_document_structure(self, text_content):
        """
        Analyze document structure to identify headers, paragraphs, lists, etc.
        
        Args:
            text_content: List of text strings from each page
            
        Returns:
            dict: Structured document information
        """
        if not self.nltk_available and not self.spacy_available:
            return {}
            
        structure = {
            "sections": [],
            "entities": [],
            "keywords": [],
            "summary": ""
        }
        
        full_text = "\n\n".join(text_content)
        
        if self.spacy_available:
            try:
                # Use SpaCy for named entity recognition
                doc = self.nlp(full_text[:1000000])  # Limit text size for performance
                
                # Extract named entities
                entities_dict = {}
                for ent in doc.ents:
                    if ent.label_ not in entities_dict:
                        entities_dict[ent.label_] = []
                    if ent.text not in entities_dict[ent.label_]:
                        entities_dict[ent.label_].append(ent.text)
                
                structure["entities"] = entities_dict
                
                # Try to identify document sections using headers
                if self.nltk_available:
                    import nltk
                    import re
                    
                    # Try to identify headers using patterns and NLP
                    lines = full_text.split('\n')
                    current_section = None
                    section_content = []
                    
                    # Common header patterns
                    header_patterns = [
                        r'^[A-Z\s]{5,}$',  # ALL CAPS LINE
                        r'^\d+\.\s+[A-Z]',  # Numbered section (1. SECTION)
                        r'^Chapter \d+',    # Chapter heading
                        r'^Section \d+',    # Section heading
                    ]
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        is_header = False
                        # Check if line matches header patterns
                        for pattern in header_patterns:
                            if re.match(pattern, line):
                                is_header = True
                                break
                        
                        # Check if this is a short line with special formatting
                        if len(line) < 80 and len(line.split()) < 10:
                            tokens = nltk.word_tokenize(line)
                            pos_tags = nltk.pos_tag(tokens)
                            
                            # Headers often start with nouns or verbs and have no punctuation
                            if pos_tags and pos_tags[0][1].startswith(('NN', 'VB')) and '.' not in line:
                                is_header = True
                        
                        if is_header:
                            # Save the previous section
                            if current_section:
                                structure["sections"].append({
                                    "title": current_section,
                                    "content": "\n".join(section_content)
                                })
                            
                            # Start a new section
                            current_section = line
                            section_content = []
                        else:
                            section_content.append(line)
                    
                    # Add the last section
                    if current_section:
                        structure["sections"].append({
                            "title": current_section,
                            "content": "\n".join(section_content)
                        })
                    
                    # If no clear sections were found, create a generic one
                    if not structure["sections"]:
                        structure["sections"].append({
                            "title": "Document Content",
                            "content": full_text
                        })
                
                # Extract keywords using frequency analysis
                import re
                from collections import Counter
                
                # Remove common stop words for better keyword extraction
                stop_words = set(self.nlp.Defaults.stop_words)
                
                # Extract potential keywords (nouns mostly)
                keywords = []
                for token in doc:
                    if (token.pos_ in ["NOUN", "PROPN"] and 
                        token.text.lower() not in stop_words and
                        len(token.text) > 2):
                        keywords.append(token.text.lower())
                
                # Count keyword frequency
                keyword_counts = Counter(keywords)
                most_common = keyword_counts.most_common(20)
                
                # Filter to meaningful keywords
                structure["keywords"] = [word for word, count in most_common if count > 1]
                
                # Generate a simple summary (first few sentences from each identified section)
                summary_parts = []
                for section in structure["sections"]:
                    content = section["content"]
                    sentences = content.split('. ')
                    
                    if sentences:
                        # Take first 2 sentences as summary for this section
                        section_summary = '. '.join(sentences[:2]) + '.'
                        summary_parts.append(section_summary)
                
                structure["summary"] = "\n\n".join(summary_parts)
                
            except Exception as e:
                logger.warning(f"Error in document structure analysis: {e}")
        
        return structure
    
    def _format_extracted_text(self, result):
        """Format extracted text with layout preservation."""
        output_lines = []
        
        # Add a header with metadata if available
        if result["metadata"]:
            output_lines.append("DOCUMENT METADATA")
            output_lines.append("=" * 50)
            for key, value in result["metadata"].items():
                output_lines.append(f"{key}: {value}")
            output_lines.append("\n")
        
        # Add extracted text
        output_lines.append("EXTRACTED TEXT")
        output_lines.append("=" * 50)
        
        for page in result["text"]:
            output_lines.append(f"\nPAGE {page['page']}\n{'-' * 30}")
            output_lines.append(page["content"])
        
        # Add structure information if available
        if result["structure"] and result["structure"].get("keywords"):
            output_lines.append("\n\nKEYWORDS")
            output_lines.append("=" * 50)
            output_lines.append(", ".join(result["structure"]["keywords"]))
        
        # Add table information if available
        if result["tables"]:
            output_lines.append("\n\nTABLES")
            output_lines.append("=" * 50)
            output_lines.append(f"The document contains {len(result['tables'])} tables.")
            output_lines.append("See the structured JSON output for table contents.")
        
        return "\n".join(output_lines)
    
    def _extract_from_excel(self, file_path):
        """Extract data from Excel files with enhanced format preservation."""
        logger.info(f"Extracting data from Excel: {file_path}")
        result = {"file_path": file_path, "type": "excel", "sheets": []}
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            # Extract file metadata
            result["metadata"] = {
                "filename": os.path.basename(file_path),
                "sheet_names": excel_file.sheet_names,
                "total_sheets": len(excel_file.sheet_names)
            }
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                sheet_info = {
                    "name": sheet_name,
                    "tables": [],
                    "charts": [],
                    "merged_cells": []
                }
                
                # Read the sheet data
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Basic sheet data
                sheet_info["rows"] = len(df)
                sheet_info["columns"] = len(df.columns)
                sheet_info["headers"] = df.columns.tolist()
                sheet_info["data"] = df.to_dict('records')
                
                # Try to extract additional metadata with openpyxl if available
                if self.excel_available:
                    try:
                        import openpyxl
                        wb = openpyxl.load_workbook(file_path, data_only=True)
                        sheet = wb[sheet_name]
                        
                        # Get merged cells
                        merged_cells = []
                        for merged_range in sheet.merged_cells.ranges:
                            merged_cells.append(str(merged_range))
                        sheet_info["merged_cells"] = merged_cells
                        
                        # Try to identify tables and named ranges
                        if hasattr(wb, 'defined_names'):
                            named_ranges = []
                            for name in wb.defined_names:
                                named_ranges.append(name.name)
                            sheet_info["named_ranges"] = named_ranges
                        
                        # Check for charts
                        charts = []
                        if hasattr(sheet, '_charts'):
                            for chart in sheet._charts:
                                chart_info = {
                                    "type": chart.type if hasattr(chart, 'type') else "unknown",
                                    "title": chart.title.text if hasattr(chart, 'title') and chart.title else "Untitled"
                                }
                                charts.append(chart_info)
                        sheet_info["charts"] = charts
                        
                    except Exception as e:
                        logger.warning(f"Error extracting advanced Excel metadata: {e}")
                
                # Save to result
                result["sheets"].append(sheet_info)
                
                # Save each sheet to CSV
                output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_{sheet_name}.csv")
                df.to_csv(output_path, index=False, encoding='utf-8')
            
            # Save structured data to JSON
            json_output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_structured.json")
            with open(json_output_path, 'w', encoding='utf-8') as f:
                json_data = {
                    "metadata": result["metadata"],
                    "sheets": []
                }
                
                # Include sheet metadata but limit the actual data to avoid huge files
                for sheet in result["sheets"]:
                    sheet_copy = sheet.copy()
                    # Only include a sample of the data (first 100 rows)
                    if "data" in sheet_copy and len(sheet_copy["data"]) > 100:
                        sheet_copy["data"] = sheet_copy["data"][:100]
                        sheet_copy["data_truncated"] = True
                    json_data["sheets"].append(sheet_copy)
                    
                json.dump(json_data, f, indent=2, ensure_ascii=False)
                
            result["structured_data_file"] = json_output_path
                
        except Exception as e:
            logger.error(f"Error extracting Excel data: {e}")
            result["error"] = str(e)
            
        return result
    
    def _extract_from_word(self, file_path):
        """Extract text and structure from Word documents with enhanced formatting."""
        logger.info(f"Extracting data from Word document: {file_path}")
        result = {
            "file_path": file_path, 
            "type": "word", 
            "text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {},
            "structure": {}
        }
        
        if not self.word_available:
            result["error"] = "python-docx library not available"
            return result
            
        try:
            doc = self.docx.Document(file_path)
            
            # Extract document properties/metadata
            try:
                core_props = doc.core_properties
                result["metadata"] = {
                    "author": core_props.author,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                    "title": core_props.title,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "comments": core_props.comments,
                    "category": core_props.category,
                    "last_modified_by": core_props.last_modified_by,
                    "revision": core_props.revision
                }
            except Exception as e:
                logger.warning(f"Error extracting document metadata: {e}")
            
            # Extract paragraphs with style information
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    para_info = {
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    }
                    
                    # Extract additional formatting information
                    if para.runs:
                        formats = []
                        for run in para.runs:
                            if run.text.strip():
                                format_info = {
                                    "text": run.text,
                                    "bold": run.bold,
                                    "italic": run.italic,
                                    "underline": run.underline,
                                    "font": run.font.name if run.font.name else None
                                }
                                formats.append(format_info)
                        if formats:
                            para_info["formats"] = formats
                            
                    paragraphs.append(para_info)
            
            result["paragraphs"] = paragraphs
            
            # Combine all text preserving paragraph breaks
            full_text = []
            for para in paragraphs:
                full_text.append(para["text"])
            result["text"] = "\n".join(full_text)
            
            # Extract tables with enhanced structure
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Extract all paragraphs from the cell
                        cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                # Get table dimensions
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                
                tables.append({
                    "id": i + 1,
                    "rows": rows,
                    "columns": cols,
                    "data": table_data
                })
            
            result["tables"] = tables
            
            # Try to determine document structure
            if self.spacy_available or self.nltk_available:
                structure = self._analyze_document_structure([result["text"]])
                result["structure"] = structure
            
            # Save extracted text to file
            output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_extracted.txt")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result["text"])
            
            # Save structured data to JSON
            json_output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_structured.json")
            structured_data = {
                "metadata": result["metadata"],
                "structure": result["structure"],
                "paragraphs": result["paragraphs"],
                "tables": result["tables"]
            }
            
            with open(json_output_path, 'w', encoding='utf-8') as f:
                json.dump(structured_data, f, indent=2, ensure_ascii=False)
                
            result["output_file"] = output_path
            result["structured_data_file"] = json_output_path
            
        except Exception as e:
            logger.error(f"Error extracting Word document data: {e}")
            result["error"] = str(e)
            
        return result
    
    def _extract_from_image(self, file_path):
        """Extract text from images using advanced OCR techniques."""
        logger.info(f"Extracting data from image: {file_path}")
        result = {
            "file_path": file_path, 
            "type": "image", 
            "text": "",
            "metadata": {},
            "layout": {}
        }
        
        if not self.ocr_available:
            result["error"] = "OCR is not available. Please check the Tesseract installation and path."
            return result
        
        try:
            # Open the image
            image = self.Image.open(file_path)
            
            # Extract image metadata
            result["metadata"] = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height
            }
            
            # Check if we need to preprocess the image
            if self.preprocess and self.cv2_available:
                # Convert to OpenCV format
                img = np.array(image)
                if len(img.shape) == 3:
                    img = self.cv2.cvtColor(img, self.cv2.COLOR_RGB2BGR)
                
                # Auto-orientation based on EXIF data
                try:
                    exif_orientation = image._getexif().get(274, 1) if hasattr(image, '_getexif') and image._getexif() else 1
                    if exif_orientation > 1:
                        orientation_map = {
                            2: (self.cv2.ROTATE_180, self.cv2.FLIP_Y),
                            3: (self.cv2.ROTATE_180, None),
                            4: (self.cv2.ROTATE_180, self.cv2.FLIP_X),
                            5: (self.cv2.ROTATE_90_COUNTERCLOCKWISE, self.cv2.FLIP_Y),
                            6: (self.cv2.ROTATE_90_CLOCKWISE, None),
                            7: (self.cv2.ROTATE_90_CLOCKWISE, self.cv2.FLIP_Y),
                            8: (self.cv2.ROTATE_90_COUNTERCLOCKWISE, None)
                        }
                        
                        if exif_orientation in orientation_map:
                            rotation, flip = orientation_map[exif_orientation]
                            if rotation:
                                img = self.cv2.rotate(img, rotation)
                            if flip:
                                img = self.cv2.flip(img, flip)
                except Exception as e:
                    logger.warning(f"Error processing image orientation: {e}")
                
                # Process the image for OCR
                # 1. Convert to grayscale if needed
                if len(img.shape) == 3:
                    gray = self.cv2.cvtColor(img, self.cv2.COLOR_BGR2GRAY)
                else:
                    gray = img
                    
                # 2. Apply noise reduction
                denoised = self.cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
                
                # 3. Apply adaptive thresholding
                binary = self.cv2.adaptiveThreshold(
                    denoised, 255, self.cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                    self.cv2.THRESH_BINARY, 11, 2
                )
                
                # 4. Apply morphological operations for better character recognition
                kernel = np.ones((1, 1), np.uint8)
                processed = self.cv2.morphologyEx(binary, self.cv2.MORPH_CLOSE, kernel)
                
                # 5. Check image size and resize if needed for better OCR
                height, width = processed.shape
                if height < 300 or width < 300 or (height > 3000 or width > 3000):
                    # Resize to a reasonable size for OCR
                    if height > width:
                        new_height = min(max(1000, height), 3000)
                        scale = new_height / height
                        new_width = int(width * scale)
                    else:
                        new_width = min(max(1000, width), 3000)
                        scale = new_width / width
                        new_height = int(height * scale)
                    
                    processed = self.cv2.resize(processed, (new_width, new_height), 
                                              interpolation=self.cv2.INTER_CUBIC)
                
                # Convert back to PIL image for Tesseract
                processed_image = self.Image.fromarray(processed)
                
                # Apply advanced OCR with specified languages
                lang_str = "+".join(self.ocr_lang)
                
                if self.preserve_layout or self.ocr_mode == "deep":
                    # Use Tesseract's layout analysis to preserve structure
                    ocr_data = self.pytesseract.image_to_data(
                        processed_image,
                        lang=lang_str,
                        output_type=self.pytesseract.Output.DICT,
                        config='--psm 1' if self.preserve_layout else self.ocr_config
                    )
                    
                    # Extract text with positional information
                    layout_text = []
                    layout_blocks = []
                    current_block = -1
                    current_line = -1
                    current_line_text = ""
                    current_block_lines = []
                    
                    for i in range(len(ocr_data["text"])):
                        text = ocr_data["text"][i].strip()
                        if not text:
                            continue
                            
                        block_num = ocr_data["block_num"][i]
                        line_num = ocr_data["line_num"][i]
                        conf = ocr_data["conf"][i]
                        
                        # Skip low confidence results
                        if conf < 30:  # Adjust confidence threshold as needed
                            continue
                            
                        # Start a new line if needed
                        if line_num != current_line:
                            if current_line_text:
                                current_block_lines.append(current_line_text)
                                current_line_text = ""
                            current_line = line_num
                        
                        # Start a new block if needed
                        if block_num != current_block:
                            if current_block_lines:
                                layout_blocks.append("\n".join(current_block_lines))
                                current_block_lines = []
                            current_block = block_num
                            
                        # Add text to current line
                        current_line_text += " " + text if current_line_text else text
                        
                    # Add the last line and block
                    if current_line_text:
                        current_block_lines.append(current_line_text)
                    if current_block_lines:
                        layout_blocks.append("\n".join(current_block_lines))
                    
                    # Save the layout information
                    result["layout"] = {
                        "blocks": layout_blocks,
                        "word_count": sum(1 for text in ocr_data["text"] if text.strip())
                    }
                    
                    # Combine all blocks for the full text
                    result["text"] = "\n\n".join(layout_blocks)
                else:
                    # Standard OCR
                    result["text"] = self.pytesseract.image_to_string(
                        processed_image, 
                        lang=lang_str,
                        config=self.ocr_config
                    )
            else:
                # Standard OCR without preprocessing
                lang_str = "+".join(self.ocr_lang)
                result["text"] = self.pytesseract.image_to_string(
                    image, 
                    lang=lang_str,
                    config=self.ocr_config
                )
            
            # Save extracted text to file
            output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_extracted.txt")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result["text"])
            
            # If layout information is available, save it to JSON
            if result["layout"]:
                json_output_path = os.path.join(self.output_dir, f"{Path(file_path).stem}_structured.json")
                structured_data = {
                    "metadata": result["metadata"],
                    "layout": result["layout"]
                }
                
                with open(json_output_path, 'w', encoding='utf-8') as f:
                    json.dump(structured_data, f, indent=2, ensure_ascii=False)
                
                result["structured_data_file"] = json_output_path
                
            result["output_file"] = output_path
            
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            traceback.print_exc()
            result["error"] = str(e)
            
        return result

def main():
    """Main function to run the data extraction tool."""
    parser = argparse.ArgumentParser(description="Advanced Data Extraction Tool")
    parser.add_argument("file_path", help="Path to the file for data extraction")
    parser.add_argument("--ocr-lang", default="eng", help="Language for OCR (default: eng). Use multiple languages with '+': eng+fra")
    parser.add_argument("--output-dir", help="Directory to save extracted data")
    parser.add_argument("--no-tables", action="store_false", dest="extract_tables",
                        help="Skip table extraction")
    parser.add_argument("--dpi", type=int, default=300, help="DPI for image conversion (default: 300)")
    parser.add_argument("--no-preprocess", action="store_false", dest="preprocess",
                        help="Skip image preprocessing")
    parser.add_argument("--no-layout", action="store_false", dest="preserve_layout",
                        help="Don't preserve document layout")
    parser.add_argument("--ocr-mode", choices=["basic", "advanced", "deep"], default="advanced",
                        help="OCR processing mode (default: advanced)")
    parser.add_argument("--no-font-analysis", action="store_false", dest="font_analysis",
                        help="Skip font analysis")
    parser.add_argument("--no-forms", action="store_false", dest="extract_forms",
                        help="Skip form field extraction")
    parser.add_argument("--threads", type=int, default=4,
                        help="Maximum number of threads for parallel processing (default: 4)")
    parser.add_argument("--tesseract-path", 
                        help="Path to Tesseract executable (e.g., 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe')")
    args = parser.parse_args()

    # Convert comma or plus separated languages to a list
    languages = args.ocr_lang.replace(',', '+').split('+')

    try:
        extractor = DataExtractor(
            ocr_lang=languages, 
            extract_tables=args.extract_tables,
            output_dir=args.output_dir,
            dpi=args.dpi,
            preprocess=args.preprocess,
            preserve_layout=args.preserve_layout,
            ocr_mode=args.ocr_mode,
            font_analysis=args.font_analysis,
            extract_forms=args.extract_forms,
            max_threads=args.threads,
            tesseract_path=args.tesseract_path
        )
        
        result = extractor.extract(args.file_path)
        
        if result and not result.get("error"):
            logger.info(f"Extraction completed successfully. Results saved to {extractor.output_dir}")
            if "output_file" in result:
                logger.info(f"Main content saved to {result['output_file']}")
            if "structured_data_file" in result:
                logger.info(f"Structured data saved to {result['structured_data_file']}")
        else:
            logger.error(f"Extraction failed: {result.get('error', 'Unknown error')}")
            sys.exit(1)
            
    except Exception as e:
        logger.error(f"An error occurred: {e}")
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()